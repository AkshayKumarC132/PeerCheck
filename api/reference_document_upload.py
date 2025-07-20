# Import necessary modules
from rest_framework.response import Response
from rest_framework.generics import CreateAPIView
from rest_framework import status
from django.db import transaction
import uuid
import os
import logging
from .models import ReferenceDocument, AuditLog, UserProfile
from .views import upload_file_to_s3, download_file_from_s3
import boto3
from peercheck import settings
from rest_framework import serializers
import requests
import tempfile
from .views import token_verification  # Assuming this is a utility function to verify tokens
from rest_framework.views import APIView
from .permissions import RoleBasedPermission  # Assuming this is a custom permission class

try:
    # AWS S3 Configuration
    S3_BUCKET_NAME = settings.AWS_STORAGE_BUCKET_NAME
    S3_REGION = settings.AWS_S3_REGION_NAME
    S3_ACCESS_KEY = settings.AWS_S3_ACCESS_KEY_ID
    S3_SECRET_KEY = settings.AWS_S3_SECRET_ACCESS_KEY

    # Initialize S3 client
    s3_client = boto3.client(
        "s3",
        region_name=S3_REGION,
        aws_access_key_id=S3_ACCESS_KEY,
        aws_secret_access_key=S3_SECRET_KEY
    )
except Exception as e:
    raise Exception(f"Failed to initialize S3 client: {str(e)}")

import PyPDF2
import pdfplumber
from docx import Document
import logging
import os
from typing import Optional

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using multiple methods for better accuracy
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If text extraction fails
    """
    try:
        text_content = ""
        
        # Method 1: Try with pdfplumber (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num} ---\n"
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    logging.info(f"Successfully extracted text from PDF using pdfplumber: {len(text_content)} characters")
                    return text_content.strip()
        except Exception as e:
            logging.warning(f"pdfplumber extraction failed: {e}")
        
        # Method 2: Fallback to PyPDF2 if pdfplumber fails
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num} ---\n"
                        text_content += page_text + "\n"
                
                if text_content.strip():
                    logging.info(f"Successfully extracted text from PDF using PyPDF2: {len(text_content)} characters")
                    return text_content.strip()
        except Exception as e:
            logging.warning(f"PyPDF2 extraction failed: {e}")
        
        # If both methods fail
        if not text_content.strip():
            raise Exception("No text could be extracted from the PDF file. The PDF might be image-based or corrupted.")
        
        return text_content.strip()
        
    except Exception as e:
        logging.error(f"PDF text extraction failed for {file_path}: {str(e)}")
        raise Exception(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX/DOC file
    
    Args:
        file_path (str): Path to the DOCX/DOC file
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If text extraction fails
    """
    try:
        text_content = ""
        
        # Handle .docx files
        if file_path.lower().endswith('.docx'):
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs_text.append(paragraph.text.strip())
            
            if paragraphs_text:
                text_content = "\n".join(paragraphs_text)
            
            # Extract text from tables if any
            tables_text = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
                
                if table_text:
                    tables_text.append("\n".join(table_text))
            
            if tables_text:
                if text_content:
                    text_content += "\n\n--- Tables ---\n\n"
                text_content += "\n\n".join(tables_text)
            
            if not text_content.strip():
                raise Exception("No text content found in the DOCX file")
            
            logging.info(f"Successfully extracted text from DOCX: {len(text_content)} characters")
            return text_content.strip()
        
        # # Handle .doc files (older format) - requires additional conversion
        # elif file_path.lower().endswith('.doc'):
        #     return extract_text_from_doc(file_path)
        
        else:
            raise Exception(f"Unsupported file format: {file_path}")
            
    except Exception as e:
        logging.error(f"DOCX text extraction failed for {file_path}: {str(e)}")
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_document(file_path: str) -> str:
    """
    Main function to extract text from various document formats
    
    Args:
        file_path (str): Path to the document file
        
    Returns:
        str: Extracted text content
    """
    if not os.path.exists(file_path):
        raise Exception(f"File does not exist: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    # elif file_extension == '.doc':
    #     return extract_text_from_doc(file_path)
    elif file_extension == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    else:
        raise Exception(f"Unsupported file format: {file_extension}")


# Enhanced text cleaning function
def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text
    
    Args:
        text (str): Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    import re
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Replace multiple newlines with double newline
    text = re.sub(r'\r\n', '\n', text)  # Normalize line endings
    text = re.sub(r'\r', '\n', text)  # Replace carriage returns
    text = re.sub(r'[ \t]+', ' ', text)  # Replace multiple spaces/tabs with single space
    
    # Remove page markers if they exist
    text = re.sub(r'\n--- Page \d+ ---\n', '\n', text)
    
    # Remove excessive blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()

class ReferenceDocumentSerializer(serializers.Serializer):
    document = serializers.FileField()
    name = serializers.CharField(required=False)
    document_type = serializers.ChoiceField(choices=['sop'], default='sop')


# 1. Upload View
class ReferenceDocumentUploadView(CreateAPIView):
    serializer_class = ReferenceDocumentSerializer
    permission_classes = [RoleBasedPermission]

    def post(self, request, token):
        user_data = token_verification(token)
        if user_data['status'] != 200:
            return Response({'error': user_data['error']}, status=status.HTTP_400_BAD_REQUEST)
        user = user_data['user']
        
        serializer = self.serializer_class(data=request.data)
        if not serializer.is_valid():
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
        
        try:
            document_file = serializer.validated_data['document']
            document_name = serializer.validated_data.get('name', document_file.name)
            document_type = serializer.validated_data.get('document_type', 'sop')
            
            allowed_extensions = ['.pdf', '.docx', '.txt']
            file_extension = os.path.splitext(document_file.name)[1].lower()

            if file_extension not in allowed_extensions:
                return Response({
                    "error": f"File type {file_extension} not supported. Allowed types: {', '.join(allowed_extensions)}"
                }, status=status.HTTP_400_BAD_REQUEST)

            with transaction.atomic():
                # Upload to S3
                s3_file_name = f"reference-documents/{uuid.uuid4()}{file_extension}"
                try:
                    file_url = upload_file_to_s3(document_file, S3_BUCKET_NAME, s3_file_name)
                    logging.info(f"Reference document uploaded to S3: {file_url}")
                except Exception as e:
                    logging.error(f"S3 upload failed: {str(e)}")
                    return Response({
                        "error": f"Failed to upload document: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                # Save initial record
                document_obj = ReferenceDocument.objects.create(
                    name=document_name,
                    document_type=document_type,
                    file_path=file_url,
                    original_filename=document_file.name,
                    file_size=document_file.size,
                    content_type=document_file.content_type,
                    upload_status='processing',
                    uploaded_by=user
                )

                # Download and extract text
                try:
                    response = requests.get(file_url, stream=True, timeout=30)
                    if response.status_code != 200:
                        raise ValueError(f"Failed to download file: HTTP {response.status_code}")

                    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
                        for chunk in response.iter_content(chunk_size=8192):
                            tmp_file.write(chunk)
                        local_file_path = tmp_file.name

                    extracted_text = self._extract_text_from_document_path(local_file_path, file_extension)

                    # Update record
                    document_obj.extracted_text = extracted_text
                    document_obj.upload_status = 'processed'
                    document_obj.save()

                    os.unlink(local_file_path)

                except Exception as e:
                    logging.error(f"Text extraction failed: {str(e)}")
                    document_obj.upload_status = 'failed'
                    document_obj.save()
                    return Response({
                        "error": f"Failed to process document: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

                # Audit log
                AuditLog.objects.create(
                    action='document_upload',
                    user=user,
                    object_id=document_obj.id,
                    object_type='ReferenceDocument',
                    details={
                        'document_name': document_name,
                        'file_size': document_file.size,
                        'document_type': document_type,
                        'file_url': file_url
                    }
                )

                return Response({
                    "document_id": document_obj.id,
                    "name": document_obj.name,
                    "document_type": document_obj.document_type,
                    "file_path": document_obj.file_path,
                    "upload_status": document_obj.upload_status,
                    "file_size": document_obj.file_size,
                    "created_at": document_obj.created_at,
                    "extracted_text_length": len(extracted_text) if extracted_text else 0
                }, status=status.HTTP_201_CREATED)

        except Exception as e:
            logging.exception("Reference document upload failed")
            return Response({
                "error": f"Document upload failed: {str(e)}"
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def _extract_text_from_document_path(self, file_path, file_extension):
        """Extract text based on file extension"""
        try:
            if file_extension == '.pdf':
                return extract_text_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return extract_text_from_docx(file_path)
            elif file_extension == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                raise ValueError(f"Unsupported file extension: {file_extension}")
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")


# 2. List View
class ReferenceDocumentListView(APIView):
    permission_classes = [RoleBasedPermission]

    def get(self, request, token):
        user_data = token_verification(token)
        if user_data['status'] != 200:
            return Response({'error': user_data['error']}, status=status.HTTP_400_BAD_REQUEST)
        user = user_data['user']
        try:
            if user:
                documents = ReferenceDocument.objects.filter(uploaded_by=user)
            else:
                documents = ReferenceDocument.objects.filter(upload_status='processed')[:50]
            documents_data = []
            for doc in documents:
                documents_data.append({
                    "id": doc.id,
                    "name": doc.name,
                    "document_type": doc.document_type,
                    "file_path": doc.file_path,
                    "upload_status": doc.upload_status,
                    "file_size": doc.file_size,
                    "original_filename": doc.original_filename,
                    "uploaded_by": doc.uploaded_by.username if doc.uploaded_by else None,
                    "created_at": doc.created_at,
                    "updated_at": doc.updated_at,
                    "has_extracted_text": bool(doc.extracted_text),
                    "related_sops_count": doc.related_sops.count()
                })
            return Response({"documents": documents_data, "total_count": len(documents_data)}, status=status.HTTP_200_OK)
        except Exception as e:
            logging.exception("Failed to list reference documents")
            return Response({"error": f"Failed to list documents: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


# 3. Detail View
class ReferenceDocumentDetailView(APIView):
    permission_classes = [RoleBasedPermission]

    def get(self, request, token, document_id):
        user_data = token_verification(token)
        if user_data['status'] != 200:
            return Response({'error': user_data['error']}, status=status.HTTP_400_BAD_REQUEST)
        
        user_role = UserProfile.objects.get(username=user_data['user']).role if user_data['user'] else None

        try:
            document = ReferenceDocument.objects.get(id=document_id)

            if user_role != 'admin':
                # Assuming 'uploaded_by' is a ForeignKey to User
                if document.uploaded_by != user_data['user']:
                    return Response({"error": "You do not have permission to view this document."},
                                    status=status.HTTP_403_FORBIDDEN)
            return Response({
                "id": document.id,
                "name": document.name,
                "document_type": document.document_type,
                "file_path": document.file_path,
                "upload_status": document.upload_status,
                "file_size": document.file_size,
                "content_type": document.content_type,
                "original_filename": document.original_filename,
                "uploaded_by": document.uploaded_by.username if document.uploaded_by else None,
                "created_at": document.created_at,
                "updated_at": document.updated_at,
                "extracted_text": document.extracted_text,
                "related_sops": [
                    {
                        "id": sop.id,
                        "name": sop.name,
                        "version": sop.version
                    } for sop in document.related_sops.all()
                ]
            }, status=status.HTTP_200_OK)
        except ReferenceDocument.DoesNotExist:
            return Response({"error": "Reference document not found"}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logging.exception("Failed to get reference document details")
            return Response({"error": f"Failed to get document details: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)