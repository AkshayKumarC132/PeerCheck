import os
import re
import string
import tempfile
import csv
import logging
import threading
import textwrap
import importlib.util
from contextlib import contextmanager, suppress
from typing import Any, Dict, List, Optional, Set, Tuple

import boto3
import numpy as np
from django.conf import settings as django_settings
from peercheck import settings
from pdf2docx import Converter
import whisper
import PyPDF2
import docx
from difflib import SequenceMatcher
from fuzzywuzzy import fuzz
from io import BytesIO
from docx.shared import RGBColor
from docx.oxml.ns import qn
import docx.oxml
import uuid
from collections import Counter, defaultdict
from sentence_transformers import SentenceTransformer, util
# Add at top (after existing ML imports like whisper/pyannote)
import torch  # If not present, add to requirements

# Detect device
def get_device():
    return torch.device("cuda" if torch.cuda.is_available() else "cpu")

DEVICE = get_device()
logging.info(f"Using device: {DEVICE}")

# Lazy global model (thread-safe)
_SENTENCE_MODEL = None
_MODEL_LOCK = threading.Lock()

def _get_sentence_model():
    global _SENTENCE_MODEL
    if _SENTENCE_MODEL is None:
        with _MODEL_LOCK:
            if _SENTENCE_MODEL is None:
                _SENTENCE_MODEL = SentenceTransformer('all-MiniLM-L6-v2', device=DEVICE)
    return _SENTENCE_MODEL

from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
import fitz  # PyMuPDF
from docx2pdf import convert

from pyannote.audio import Pipeline, Inference, Model
from pyannote.core import Segment

from .models import SpeakerProfile
from .speaker_utils import match_speaker_embedding

# Load Whisper model once
model = whisper.load_model(getattr(settings, 'WHISPER_MODEL', 'small.en'), device=DEVICE)

_DIARIZATION_PIPELINE: Optional[Pipeline] = None
_EMBEDDING_INFERENCE: Optional[Inference] = None
_MODEL_LOCK = threading.Lock()


def _get_hf_token() -> Optional[str]:
    """Return the Hugging Face token if available."""
    token = getattr(settings, "HF_TOKEN", None) or getattr(django_settings, "HF_TOKEN", None)
    return token


def _get_diarization_pipeline() -> Pipeline:
    """Lazily initialise the diarization pipeline."""
    global _DIARIZATION_PIPELINE
    if _DIARIZATION_PIPELINE is None:
        with _MODEL_LOCK:
            if _DIARIZATION_PIPELINE is None:
                _DIARIZATION_PIPELINE = Pipeline.from_pretrained(
                    "pyannote/speaker-diarization-3.1",
                    use_auth_token=_get_hf_token(),
                )
                _DIARIZATION_PIPELINE.to(DEVICE)
    return _DIARIZATION_PIPELINE


def _get_embedding_inference() -> Inference:
    """Lazily initialise the speaker embedding inference model."""
    global _EMBEDDING_INFERENCE
    if _EMBEDDING_INFERENCE is None:
        with _MODEL_LOCK:
            if _EMBEDDING_INFERENCE is None:
                embedding_model = Model.from_pretrained(
                    "pyannote/embedding",
                    use_auth_token=_get_hf_token(),
                )
                _EMBEDDING_INFERENCE = Inference(embedding_model, window="whole", device=DEVICE)
    return _EMBEDDING_INFERENCE

# Initialize S3 client
s3_client = boto3.client(
    's3',
    aws_access_key_id=settings.AWS_S3_ACCESS_KEY_ID,
    aws_secret_access_key=settings.AWS_S3_SECRET_ACCESS_KEY,
    region_name=settings.AWS_S3_REGION_NAME
)

def allowed_file(filename, allowed_extensions):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def upload_file_to_s3(file_obj, s3_key):
    print("Upload File to S3 ", file_obj, s3_key)
    """Upload file to S3 and return the S3 URL"""
    try:
        s3_client.upload_fileobj(
            file_obj,
            settings.AWS_STORAGE_BUCKET_NAME,
            s3_key,
            ExtraArgs={'ACL': 'private'}
        )
        return f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.amazonaws.com/{s3_key}"
    except Exception as e:
        raise Exception(f"Failed to upload to S3: {str(e)}")

def download_file_from_s3(s3_key):
    """Download file from S3 to temporary file and return path"""
    try:
        response = s3_client.get_object(
            Bucket=settings.AWS_STORAGE_BUCKET_NAME,
            Key=s3_key
        )

        _, ext = os.path.splitext(s3_key)
        suffix = ext if ext else None

        # Create temporary file while preserving the original extension so
        # downstream consumers can perform extension-based handling (for
        # example, DOCX conversion before PDF highlighting).
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
        temp_file.write(response['Body'].read())
        temp_file.close()

        return temp_file.name
    except Exception as e:
        raise Exception(f"Failed to download from S3: {str(e)}")

def get_s3_key_from_url(s3_url):
    """Extract S3 key from S3 URL"""
    return s3_url.replace(f"https://{settings.AWS_STORAGE_BUCKET_NAME}.s3.amazonaws.com/", "")

def extract_text_txt(file_path):
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

def extract_text_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ''
    return text

def extract_text_docx(file_path):
    document = docx.Document(file_path)
    return '\n'.join([para.text for para in document.paragraphs])

def extract_text_from_s3(s3_url):
    """Extract text from document stored in S3"""
    s3_key = get_s3_key_from_url(s3_url)
    temp_path = download_file_from_s3(s3_key)
    
    try:
        ext = s3_key.rsplit('.', 1)[1].lower()
        if ext == 'pdf':
            text = extract_text_pdf(temp_path)
        elif ext == 'docx':
            text = extract_text_docx(temp_path)
        elif ext == 'txt':
            text = extract_text_txt(temp_path)
        else:
            text = ''
        return text
    finally:
        # Clean up temp file
        os.unlink(temp_path)

def transcribe_audio_from_s3(s3_url):
    """Transcribe audio file stored in S3 and return full Whisper result (with word-level timestamps)"""
    s3_key = get_s3_key_from_url(s3_url)
    temp_path = download_file_from_s3(s3_key)
    try:
        # Use word_timestamps=True to get word-level info
        result = model.transcribe(temp_path, word_timestamps=True)
        if not result or not result.get("text"):
            raise ValueError("Transcription result is empty or invalid.")
        return result  # Return the full result dict
    except Exception as e:
        logging.error(f"Failed to transcribe audio from S3: {e}")
        raise
    finally:
        os.unlink(temp_path)

def normalize_line(s: str) -> str:
    s = s.lower()
    s = re.sub(r"[\[\]\(\)\{\}\<\>]", "", s)
    s = s.translate(str.maketrans('', '', string.punctuation))
    return ' '.join(s.split())

def find_missing(text, transcript, threshold=0.6):
    """
    Fuzzy-match each original text line against all transcript lines.
    Returns: (matched_html, missing_html, matched_words, total_words, entire_html)
    """
    text_lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    trans_lines = [ln.strip() for ln in transcript.splitlines() if ln.strip()]
    norm_text = [normalize_line(ln) for ln in text_lines]
    norm_trans = [normalize_line(ln) for ln in trans_lines]

    # Detect start offset
    start_idx = 0
    if norm_trans and norm_text:
        K = min(len(norm_trans), 3)
        best_avg, best_idx = 0.0, 0
        for i in range(len(norm_text) - K + 1):
            avg = sum(
                SequenceMatcher(None, norm_text[i + j], norm_trans[j]).ratio()
                for j in range(K)
            ) / K
            if avg > best_avg:
                best_avg, best_idx = avg, i
        if best_avg >= threshold:
            start_idx = best_idx

    matched_spans, missing_spans = [], []
    matched_words = 0
    total_words = sum(len(normalize_line(ln).split()) for ln in text_lines)

    # Mark lines before start offset as missing
    for ln in text_lines[:start_idx]:
        missing_spans.append(f'<span class="missing">{ln}</span>')

    # Fuzzy-match remaining lines
    for orig, norm in zip(text_lines[start_idx:], norm_text[start_idx:]):
        best = max(
            (fuzz.token_set_ratio(norm, tnorm) for tnorm in norm_trans),
            default=0
        ) / 100.0
        wcount = len(norm.split())
        if best >= threshold:
            matched_spans.append(f'<span class="match">{orig}</span>')
            matched_words += wcount
        else:
            missing_spans.append(f'<span class="missing">{orig}</span>')

    m_html = '<p>' + '</p><p>'.join(matched_spans) + '</p>' if matched_spans else ''
    x_html = '<p>' + '</p><p>'.join(missing_spans) + '</p>' if missing_spans else ''
    
    # Build entire document view
    entire_spans = []
    for orig, norm in zip(text_lines, norm_text):
        best = max((fuzz.token_set_ratio(norm, tnorm) for tnorm in norm_trans), default=0) / 100.0
        cls = "match" if best >= threshold else "missing"
        entire_spans.append(f'<span class="{cls}">{orig}</span>')
    entire_html = '<p>' + '</p><p>'.join(entire_spans) + '</p>'
    
    return m_html, x_html, matched_words, total_words, entire_html

def highlight_docx_cross_platform(docx_path, norm_trans, output_path, threshold=0.6):
    """
    Highlights text in a DOCX file using python-docx (cross-platform).

    Args:
        docx_path (str): Path to the input DOCX file.
        norm_trans (list): A list of normalized transcript strings.
        output_path (str): Path to save the highlighted DOCX.
        threshold (float): Similarity threshold for highlighting.
    """
    document = docx.Document(docx_path)
    
    # Define colors
    GREEN = RGBColor(0, 176, 80)
    RED = RGBColor(255, 0, 0)

    # 1. Highlight regular paragraphs
    for para in document.paragraphs:
        if not para.text.strip():
            continue
        norm_para_text = normalize_line(para.text)
        best_score = max((fuzz.token_set_ratio(norm_para_text, t) for t in norm_trans), default=0) / 100.0
        color = GREEN if best_score >= threshold else RED
        for run in para.runs:
            run.font.color.rgb = color

    # 2. Highlight text within shapes and text boxes by manipulating the underlying XML
    tree = document.element.body
    text_runs_in_shapes = tree.xpath('.//w:txbxContent//w:t') # Find all text runs in text boxes

    for text_run in text_runs_in_shapes:
        parent_paragraph = text_run.getparent().getparent()
        texts_in_p = parent_paragraph.xpath('.//w:t/text()')
        full_text = "".join(texts_in_p).strip()

        if not full_text:
            continue
        
        norm_shape_text = normalize_line(full_text)
        best_score = max((fuzz.token_set_ratio(norm_shape_text, t) for t in norm_trans), default=0) / 100.0
        color = GREEN if best_score >= threshold else RED
        
        # Apply color by creating/updating XML properties
        run_properties = text_run.getparent().find(docx.oxml.ns.qn('w:rPr'))
        if run_properties is None:
            run_properties = docx.oxml.OxmlElement('w:rPr')
            text_run.getparent().insert(0, run_properties)
        
        color_element = run_properties.find(docx.oxml.ns.qn('w:color'))
        if color_element is None:
            color_element = docx.oxml.OxmlElement('w:color')
            run_properties.append(color_element)
            
        color_element.set(docx.oxml.ns.qn('w:val'), str(color))

    document.save(output_path)


def _load_abbreviation_map(csv_path: str) -> dict:
    """Load abbreviation mappings from a CSV file.

    Supports rows that contain multiple abbreviations for the same meaning
    separated by characters such as commas, slashes, dashes, or spaces. Each
    abbreviation variant is normalized by stripping non-alphanumeric
    characters and upper-casing before being added to the map. Single-character
    entries and headers are skipped.
    """
    abbr_map: dict[str, dict] = {}
    for enc in ("utf-8", "cp1252", "latin1"):
        try:
            with open(csv_path, newline="", encoding=enc) as csvfile:
                reader = csv.reader(csvfile)
                for row in reader:
                    if len(row) < 2:
                        continue
                    abbr_field, full = row[0].strip(), row[1].strip()
                    if not (abbr_field and full) or abbr_field.upper() == "ABBREVIATION":
                        continue

                    # Split the abbreviation field into individual candidates
                    for part in re.split(r",", abbr_field):
                        part = part.strip()
                        if not part:
                            continue

                        # Consider sub-variants separated by '/', '-' or spaces
                        candidates = {part}
                        for sep in ("/", "-", " "):
                            if sep in part:
                                candidates.update(
                                    seg.strip() for seg in part.split(sep) if seg.strip()
                                )

                        for abbr in candidates:
                            norm = re.sub(r"[^A-Za-z0-9]", "", abbr).upper()
                            if len(norm) < 2:
                                continue
                            entry = abbr_map.setdefault(norm, {"abbrs": set(), "full": full})
                            entry["abbrs"].add(abbr)

            for data in abbr_map.values():
                data["abbrs"] = list(data["abbrs"])
            logging.info(
                "Loaded %d abbreviations from CSV using %s", len(abbr_map), enc
            )
            return abbr_map
        except Exception as exc:
            logging.warning("Failed to load acronyms CSV with %s: %s", enc, exc)
    return abbr_map


def _confirm_abbreviations(abbr_map: dict, transcript: str) -> dict:
    """Filter abbreviation map to only those whose full form appears in transcript."""
    transcript_lower = transcript.lower()
    validated = {}
    for norm, data in abbr_map.items():
        full = data["full"].lower()
        abbr_display = "/".join(data.get("abbrs", []))
        if full in transcript_lower:
            validated[norm] = data
            logging.debug(
                "Exact transcript match for abbreviation '%s' -> '%s'",
                abbr_display,
                data["full"],
            )
            continue
        ratio = fuzz.token_set_ratio(full, transcript_lower)
        if ratio >= 65:
            validated[norm] = data
            logging.debug(
                "Fuzzy transcript match for abbreviation '%s' -> '%s' with ratio %s",
                abbr_display,
                data["full"],
                ratio,
            )
    logging.info("Transcript validation enabled: %d abbreviations confirmed", len(validated))
    return validated


def _replace_validated_abbreviations(doc, validated_abbrs: dict) -> None:
    """Replace red highlights with dark green for validated abbreviations."""
    if not validated_abbrs:
        logging.info("No validated abbreviations to apply")
        return

    dark_green = (0.6, 1, 0.6)
    light_red = (1, 0.6, 0.6)

    def _is_light_red(color, target=light_red, tol=0.1):
        return all(abs(color[i] - target[i]) <= tol for i in range(3))

    for page_num in range(len(doc)):
        page = doc.load_page(page_num)

        red_annots = []
        annot = page.first_annot
        while annot:
            colors = annot.colors or {}
            stroke = colors.get("stroke")
            if stroke and _is_light_red(stroke):
                red_annots.append((annot, fitz.Rect(annot.rect)))
            annot = annot.next

        if not red_annots:
            continue

        words_on_page = page.get_text("words")
        for w in words_on_page:
            word_text = w[4].strip()
            norm_word = re.sub(r"[^A-Za-z0-9]", "", word_text).upper()
            data = validated_abbrs.get(norm_word)
            if not data:
                continue
            rect = fitz.Rect(w[:4])
            for annot, a_rect in list(red_annots):
                if a_rect.intersects(rect):
                    page.delete_annot(annot)
                    red_annots.remove((annot, a_rect))
                    new_annot = page.add_highlight_annot(rect)
                    new_annot.set_colors(stroke=dark_green)
                    new_annot.set_opacity(0.6)
                    new_annot.update()
                    logging.info(
                        "Validated and highlighted abbreviation '%s' on page %d",
                        word_text,
                        page_num + 1,
                    )
                    break

@contextmanager
def _word_com_apartment() -> Any:
    """Initialise a COM apartment suitable for Microsoft Word automation."""

    spec = importlib.util.find_spec("pythoncom")
    if spec is None:
        raise ValueError(
            "Converting Office documents to PDF requires the 'pywin32' package "
            "and a Microsoft Word installation."
        )

    import pythoncom  # type: ignore

    com_initialized = False
    try:
        try:
            pythoncom.CoInitialize()
            com_initialized = True
        except Exception as exc:  # pragma: no cover - best effort initialisation
            # RPC_E_CHANGED_MODE occurs when the thread is already initialised
            # in a different COM apartment. In that case we fall back to
            # explicitly initialising an STA apartment which Word automation
            # requires.
            if getattr(exc, "hresult", None) == -2147417850:
                pythoncom.CoInitializeEx(pythoncom.COINIT_APARTMENTTHREADED)
                com_initialized = True
            else:
                raise ValueError(
                    "Failed to initialise COM before launching Word: {0}".format(exc)
                ) from exc

        yield pythoncom
    finally:
        if com_initialized:
            pythoncom.CoUninitialize()


def _convert_office_document_to_pdf(source_path: str, output_path: str) -> None:
    """Convert DOCX or DOC files to PDF using the best available backend."""

    extension = os.path.splitext(source_path)[1].lower()
    if extension == ".docx":
        try:
            with _word_com_apartment():
                convert(source_path, output_path)
        except Exception as exc:
            raise ValueError(
                f"Failed to convert DOCX file '{source_path}' to PDF using Word: {exc}"
            ) from exc
        return

    if extension == ".doc":
        if importlib.util.find_spec("win32com.client") is None:
            raise ValueError(
                "Converting legacy .doc files to PDF requires the 'pywin32' package "
                "and a Microsoft Word installation."
            )

        from win32com.client import DispatchEx  # type: ignore

        with _word_com_apartment():
            try:
                word = DispatchEx("Word.Application")
            except Exception as exc:
                raise ValueError(
                    f"Failed to launch Microsoft Word for DOC conversion: {exc}"
                ) from exc

            word.Visible = False
            try:
                try:
                    document = word.Documents.Open(os.path.abspath(source_path))
                except Exception as exc:
                    raise ValueError(
                        f"Failed to open DOC file '{source_path}' with Word: {exc}"
                    ) from exc

                try:
                    document.SaveAs(os.path.abspath(output_path), FileFormat=17)
                except Exception as exc:
                    raise ValueError(
                        f"Failed to export DOC file '{source_path}' to PDF using Word: {exc}"
                    ) from exc
                finally:
                    document.Close(False)
            finally:
                word.Quit()

        return

    raise ValueError(
        f"Unsupported Office document extension '{extension}' for PDF conversion."
    )


def generate_highlighted_pdf(doc_path, query_text, output_path, require_transcript_match=True):
    """
    Opens a document, identifies relevant pages, highlights text on those pages based on semantic and numeric matching,
    and saves the result to a new PDF file.
    Adds robust validation for PDF input.
    """
    import logging
    pdf_path = doc_path
    source_was_office_doc = doc_path.lower().endswith(('.docx', '.doc'))
    if source_was_office_doc:
        temp_pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        try:
            _convert_office_document_to_pdf(doc_path, temp_pdf_path)
        except Exception as exc:
            if os.path.exists(temp_pdf_path):
                os.unlink(temp_pdf_path)
            raise ValueError(
                f"Failed to convert Office document '{doc_path}' to PDF: {exc}"
            )
        pdf_path = temp_pdf_path

    # --- Robust PDF Validation ---
    if not os.path.exists(pdf_path) or os.path.getsize(pdf_path) == 0:
        raise ValueError(f"Input file '{pdf_path}' does not exist or is empty.")
    with open(pdf_path, 'rb') as f:
        header = f.read(5)
        if header != b'%PDF-':
            raise ValueError(f"Input file '{pdf_path}' is not a valid PDF (missing %PDF header).")
    try:
        target_doc = fitz.open(pdf_path)
    except Exception as e:
        logging.error(f"Failed to open PDF '{pdf_path}': {e}")
        raise ValueError(f"Failed to open PDF '{pdf_path}': {e}")
    
    # --- Helper Functions ---
    def split_into_sentences(text):
        return [s.strip() for s in re.split(r'(?<=[.?!])\s+', text) if len(s.strip()) > 5]

    def extract_numeric_patterns(text):
        return set(re.findall(r'\b(?:\d+(?:\.\d+)*|\w+-\w+|\w{3,})\b', text.lower()))

    # --- 1. Page Relevance Scoring ---
    from rapidfuzz import fuzz

    query_tokens = extract_numeric_patterns(query_text)
    
    token_freq = Counter()
    target_page_tokens = []
    for page in target_doc:
        page_text = page.get_text()
        tokens = extract_numeric_patterns(page_text)
        target_page_tokens.append(tokens)
        token_freq.update(tokens)

    COMMON_TOKEN_THRESHOLD = int(0.5 * len(target_doc))
    filtered_query_tokens = {tok for tok in query_tokens if token_freq[tok] < COMMON_TOKEN_THRESHOLD}

    def token_score(token):
        return len(token) / (1 + token_freq[token])

    def fuzzy_match_score(query_set, target_set):
        score = 0
        for q in query_set:
            for t in target_set:
                if fuzz.ratio(q, t) > 90:
                    score += token_score(q)
        return score

    page_scores = []
    for i, tgt_tokens in enumerate(target_page_tokens):
        score = fuzzy_match_score(filtered_query_tokens, tgt_tokens)
        page_scores.append((i, score))

    page_scores = sorted(page_scores, key=lambda x: x[1], reverse=True)

    top_k = 5
    top_pages = set()
    for i, score in page_scores[:top_k]:
        top_pages.update([i - 1, i, i + 1])
    relevant_page_nums = sorted([i for i in top_pages if 0 <= i < len(target_doc)])
    
    # --- 2. Semantic Matching Setup ---
    model = SentenceTransformer('all-MiniLM-L6-v2', device=DEVICE)
    query_chunks = split_into_sentences(query_text)
    if not query_chunks: query_chunks = [query_text]
    query_embeddings = model.encode(query_chunks, convert_to_tensor=True)

    # --- 3. Apply Highlighting to Relevant Pages ---
    for page_num in relevant_page_nums:
        page = target_doc.load_page(page_num)
        page_text = page.get_text("text")
        page_chunks = split_into_sentences(page_text)
        if not page_chunks:
            page_chunks = [page_text]

        page_embeddings = model.encode(page_chunks, convert_to_tensor=True)

        cosine_scores = util.cos_sim(page_embeddings, query_embeddings)

        matched_page_chunks = set()
        import torch
        for i, row in enumerate(cosine_scores):
            if len(row) > 0 and torch.max(row) > 0.7:
                if i < len(page_chunks):
                    matched_page_chunks.add(page_chunks[i].lower())

        green_rects = []
        words_on_page = page.get_text("words")
        for w in words_on_page:
            word_text = w[4].strip()
            word_text_lower = word_text.lower()
            rect = fitz.Rect(w[:4])

            is_numeric_match = any(
                word_text_lower == t or word_text_lower in t or t in word_text_lower
                for t in filtered_query_tokens
            )
            is_semantic_match = any(word_text_lower in chunk for chunk in matched_page_chunks)

            # Define light green and light red colors
            light_green = (0.6, 1, 0.6)
            light_red = (1, 0.6, 0.6)

            if is_numeric_match or is_semantic_match:
                highlight = page.add_highlight_annot(rect)
                highlight.set_colors(stroke=light_green)
                highlight.set_opacity(0.3)
                highlight.update()
                green_rects.append(rect)
            else:
                if not any(rect.intersects(g) for g in green_rects):
                    highlight = page.add_highlight_annot(rect)
                    highlight.set_colors(stroke=light_red)
                    highlight.set_opacity(0.3)
                    highlight.update()

    # --- Post-highlight Cleanup Pass for Abbreviations ---
    logging.info("Starting abbreviation cleanup pass")
    csv_path = os.path.join(settings.BASE_DIR, "Acronyms1.csv")
    abbreviations = _load_abbreviation_map(csv_path)

    if require_transcript_match:
        validated_abbrs = _confirm_abbreviations(abbreviations, query_text)
    else:
        validated_abbrs = abbreviations
        logging.info(
            "Transcript validation disabled: using %d abbreviations directly",
            len(validated_abbrs),
        )

    _replace_validated_abbreviations(target_doc, validated_abbrs)
    # --- 4. Save the Output ---
    try:
        target_doc.save(output_path, garbage=4, deflate=True)
    finally:
        target_doc.close()
        if source_was_office_doc and os.path.exists(pdf_path):
            try:
                os.unlink(pdf_path)
            except PermissionError:
                logging.warning(
                    "Temp file in use after Office conversion, skipping deletion: %s",
                    pdf_path,
                )
            except OSError as exc:
                logging.warning(
                    "Failed to delete temporary Office conversion file %s: %s",
                    pdf_path,
                    exc,
                )

    return output_path

def create_highlighted_pdf_document(
    text_s3_url,
    transcript,
    require_transcript_match=True,
    three_pc_entries: Optional[List[Dict]] = None,
    use_llm: bool = True,
    llm_provider: str = "ollama",
):
    """
    Orchestrates the generation of a highlighted PDF using LLM analysis.
    """
    output_filename = f"processed/{uuid.uuid4()}_highlighted_report.pdf"

    # Download the reference document
    s3_key = get_s3_key_from_url(text_s3_url)
    temp_input_path = download_file_from_s3(s3_key)

    temp_output_path = ""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_output_path = temp_file.name

        interactions: List[Dict[str, Any]] = []

        if use_llm:
            # --- KEY CHANGE: LLM does the analysis inside generate_highlighted_pdf ---
            # It returns the structured analysis (including interactions)
            llm_analysis = generate_highlighted_pdf_with_llm(
                temp_input_path,
                transcript,
                temp_output_path,
                provider=llm_provider,
            )

            # Use LLM-detected interactions if available, otherwise fallback (or empty)
            interactions = llm_analysis.get("interactions", []) if llm_analysis else []
        else:
            # Use semantic/keyword highlighting without an LLM
            generate_highlighted_pdf(
                temp_input_path,
                transcript,
                temp_output_path,
                require_transcript_match=require_transcript_match,
            )

        # If user explicitly passed entries (rare now), use them, else use LLM's
        final_entries = three_pc_entries if three_pc_entries is not None else interactions

        if final_entries:
            # Map LLM "interactions" to the format expected by summary table if needed
            # The LLM output keys (speaker, text, role, status) match our needs
            append_three_pc_summary_to_pdf(temp_output_path, final_entries)

        with open(temp_output_path, 'rb') as f_out:
            output_s3_url = upload_file_to_s3(f_out, output_filename)

        return output_s3_url

    finally:
        for path in (temp_input_path, temp_output_path):
            if not path:
                continue
            try:
                if os.path.exists(path):
                    os.unlink(path)
            except PermissionError:
                logging.warning(f"Temp file in use, skipping deletion: {path}")

def _locate_quote(doc, quote: str, search_last: bool = False):
    if not quote:
        return None

    found = None
    for p_idx, page in enumerate(doc):
        hits = page.search_for(quote)
        if not hits:
            continue
        if search_last:
            found = (p_idx, hits[-1])
        else:
            return (p_idx, hits[0])

    return found


def generate_highlighted_pdf_with_llm(doc_path, transcript_text, output_path, provider: str = "ollama"):
    """
    Use the Local LLM to analyze the transcript + document and:
      1. Identify the relevant section number from the procedure and its start/end quotes.
      2. Within that *Active Region* only, apply word-level Green/Red highlighting:
         - Green: word appears in the transcript (spoken).
         - Red:   word does not appear in the transcript (missed/skipped).
      3. Leave all text outside the Active Region uncolored.
    """
    from .llm_service import analyze_3pc
    import fitz

    # 1. Load PDF and Extract Text
    try:
        doc = fitz.open(doc_path)
    except Exception as e:
        raise ValueError(f"Failed to open PDF: {e}")

    full_pdf_text = ""
    for page in doc:
        full_pdf_text += page.get_text()

    # 2. Analyze with Local LLM
    logging.info("Sending Transcript + Procedure to Local LLM for analysis (3PC + section detection)...")
    analysis = analyze_3pc(transcript_text, full_pdf_text, provider=provider)
    
    start_quote = (analysis.get("highlight_start_quote") or "").strip()
    end_quote = (analysis.get("highlight_end_quote") or "").strip()
    section_id = (analysis.get("relevant_section_number") or "").strip() or "Unknown"

    logging.info("LLM identified Section %s", section_id)
    logging.info("Highlighting region from '%s' to '%s'", start_quote, end_quote)

    # Pre-compute normalized transcript word set for Green/Red logic
    def _norm_word(w: str) -> str:
        return re.sub(r"[^A-Za-z0-9]", "", w).lower()

    transcript_word_set = {
        _norm_word(w) for w in transcript_text.split() if _norm_word(w)
    }

    # 3. Define Active Region using section-aware ranges (supports multiple mentions)
    GREEN = (0.6, 1.0, 0.6)
    RED = (1.0, 0.6, 0.6)

    # --- LOAD ACRONYM MAPPINGS ---
    # Maps acronym (uppercase) -> list of spoken words (lowercase)
    # e.g., "AL" -> ["alpha", "lima"]
    acronym_to_spoken = {}
    try:
        import csv
        csv_path = os.path.join(settings.BASE_DIR, "Acronyms1.csv")
        with open(csv_path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            for row in reader:
                abbr = row.get("Abbreviation", "").strip().upper()
                meaning = row.get("Meaning", "").strip()
                if abbr and meaning:
                    # Convert meaning to individual words (lowercase)
                    spoken_words = [_norm_word(w) for w in meaning.split() if _norm_word(w)]
                    if spoken_words:
                        acronym_to_spoken[abbr] = spoken_words
        logging.info(f"Loaded {len(acronym_to_spoken)} acronym mappings for word matching")
    except Exception as e:
        logging.warning(f"Could not load Acronyms1.csv: {e}")

    def word_matches_transcript(pdf_word: str) -> bool:
        """
        Returns True if the PDF word appears in the transcript OR
        if any word from the acronym's expansion appears in transcript.
        """
        norm = _norm_word(pdf_word)
        if not norm:
            return False

        # Direct match
        if norm in transcript_word_set:
            return True

        # Acronym expansion match
        # Check if the PDF word (uppercase) is an acronym
        upper_word = pdf_word.strip().upper()
        if upper_word in acronym_to_spoken:
            # Check if ANY of the spoken expansion words appear in transcript
            spoken_words = acronym_to_spoken[upper_word]
            for spoken_w in spoken_words:
                if spoken_w in transcript_word_set:
                    return True

        return False

    def _single_region(start_loc, end_loc):
        """Build a single active region spanning from start to end across pages."""

        if not start_loc:
            return {}

        # Default the end to the final page if the quote is missing
        if not end_loc:
            end_page = doc.page_count - 1
            end_rect = doc[end_page].rect
            end_loc = (end_page, end_rect)

        start_page, start_rect = start_loc
        end_page, end_rect = end_loc

        if start_page > end_page:
            end_page = start_page
            end_rect = start_rect

        regions: Dict[int, List[fitz.Rect]] = defaultdict(list)

        for p_idx in range(start_page, end_page + 1):
            page = doc[p_idx]
            page_rect = page.rect

            if start_page == end_page:
                active_rect = fitz.Rect(
                    page_rect.x0,
                    start_rect.y0,
                    page_rect.x1,
                    end_rect.y1,
                )
            elif p_idx == start_page:
                active_rect = fitz.Rect(
                    page_rect.x0,
                    start_rect.y0,
                    page_rect.x1,
                    page_rect.y1,
                )
            elif p_idx == end_page:
                active_rect = fitz.Rect(
                    page_rect.x0,
                    page_rect.y0,
                    page_rect.x1,
                    end_rect.y1,
                )
            else:
                active_rect = page_rect

            regions[p_idx].append(active_rect)

        return regions

    # Locate the active region using the available anchors (single section only)
    page_regions: Dict[int, List[fitz.Rect]] = defaultdict(list)

    if start_quote or end_quote or section_id != "Unknown":
        start_loc = _locate_quote(doc, start_quote or section_id)

        if not start_loc and section_id and section_id != "Unknown":
            start_loc = _locate_quote(doc, section_id)

        end_loc = _locate_quote(doc, end_quote or start_quote or section_id, search_last=True)

        if not start_loc:
            logging.warning(
                "Unable to locate highlight anchors (start='%s', section='%s'). Skipping section-scoped highlighting.",
                start_quote,
                section_id,
            )
        else:
            if not end_loc:
                logging.warning(
                    "LLM returned incomplete highlight anchors (start='%s', end='%s'). Using available text for a single-region highlighting.",
                    start_quote,
                    end_quote,
                )

            for p_idx, rects in _single_region(start_loc, end_loc).items():
                page_regions[p_idx].extend(rects)
    else:
        logging.warning(
            "LLM did not return highlight anchors. Skipping section-scoped highlighting."
        )

    if page_regions:
        for p_idx, regions in page_regions.items():
            page = doc[p_idx]
            words_on_page = page.get_text("words")

            for w in words_on_page:
                x0, y0, x1, y1, word_text, *_ = w
                rect = fitz.Rect(x0, y0, x1, y1)

                if not any(region.intersects(rect) for region in regions):
                    continue

                norm = _norm_word(word_text)

                if not norm:
                    continue

                color = GREEN if word_matches_transcript(word_text) else RED

                annot = page.add_highlight_annot(rect)
                annot.set_colors(stroke=color)
                annot.set_opacity(0.3)
                annot.update()
    # Save
    doc.save(output_path)
    doc.close()
    
    return analysis

def create_highlighted_docx_from_s3(
    text_s3_url,
    transcript,
    high_threshold=0.6,
    low_threshold=0.3,
    three_pc_entries: Optional[List[Dict]] = None,
):
    """
    Generates a highlighted DOCX report from a reference document (PDF or DOCX)
    and a transcript, then uploads it to S3.
    """
    # Use dummy S3 functions for local paths if not using actual S3
    # Detect if the input is an S3 URL (either s3:// or https://...amazonaws.com/)
    is_s3_url = text_s3_url.startswith("s3://") or (
        text_s3_url.startswith("https://") and ".amazonaws.com/" in text_s3_url
    )
    s3_key = get_s3_key_from_url(text_s3_url) if is_s3_url else text_s3_url
    
    # Always download if it's an S3 URL (s3:// or https://...amazonaws.com/)
    temp_input_path = download_file_from_s3(s3_key) if is_s3_url else s3_key
    
    docx_in_path = None
    output_path = None
    
    try:
        norm_trans = [normalize_line(ln) for ln in transcript.splitlines() if ln.strip()]
        ext = s3_key.rsplit('.', 1)[-1].lower()
        
        # --- CONVERT INPUT FILE TO DOCX IF NECESSARY ---
        if ext == 'docx':
            print("Input is a DOCX file. No conversion needed.")
            docx_in_path = temp_input_path
        elif ext == 'pdf':
            print("Input is a PDF file. Converting to DOCX...")
            docx_in_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
            cv = Converter(temp_input_path)
            cv.convert(docx_in_path, start=0, end=None)
            cv.close()
            print(f"Conversion complete. Temporary DOCX at: {docx_in_path}")
        else:
            raise ValueError(f"Unsupported file type for reference document: '{ext}'")

        # --- HIGHLIGHT THE DOCX AND CREATE THE FINAL REPORT ---
        output_path = tempfile.NamedTemporaryFile(delete=False, suffix='.docx').name
        
        highlight_docx_three_color(
            docx_path=docx_in_path,
            norm_trans=norm_trans,
            output_path=output_path,
            high_threshold=high_threshold,
            low_threshold=low_threshold
        )

        if three_pc_entries is not None:
            append_three_pc_summary_to_docx(output_path, three_pc_entries)
        
        # --- UPLOAD THE FINAL REPORT TO S3 ---
        output_filename = os.path.basename(s3_key).rsplit('.', 1)[0]
        output_s3_key = f"processed/{uuid.uuid4()}_{output_filename}.docx"
        
        with open(output_path, 'rb') as f:
            output_s3_url = upload_file_to_s3(f, output_s3_key) if is_s3_url else output_path
        
        return output_s3_url
        
    finally:
        # --- Clean up all temporary files ---
        if is_s3_url and temp_input_path and os.path.exists(temp_input_path):
            os.unlink(temp_input_path)
        if docx_in_path and docx_in_path != temp_input_path and os.path.exists(docx_in_path):
            os.unlink(docx_in_path)
        if output_path and os.path.exists(output_path) and is_s3_url: # Keep local file if not using S3
            os.unlink(output_path)

def _is_other_speaker(name: Optional[str]) -> bool:
    if not name:
        return False
    return name.strip().lower() in {"other", "other speaker"}


def _drop_overlapping_duplicates(segments: List[Dict], min_overlap: float = 0.25) -> List[Dict]:
    """Remove overlapping diarization segments that are likely duplicates.

    The filter is conservative: it keeps longer, higher-content segments and
    prefers non-placeholder speakers when two segments substantially overlap or
    touch with nearly identical text.
    """
    if not segments:
        return []

    cleaned: List[Dict] = []

    def _duration(seg: Dict) -> float:
        try:
            return float(seg.get("duration") or (seg.get("end") - seg.get("start")))
        except Exception:
            return 0.0

    segments = sorted(
        segments, key=lambda s: (float(s.get("start", 0.0)), float(s.get("end", 0.0)))
    )

    for seg in segments:
        if not cleaned:
            cleaned.append(seg)
            continue

        prev = cleaned[-1]
        try:
            prev_start = float(prev.get("start"))
            prev_end = float(prev.get("end"))
            seg_start = float(seg.get("start"))
            seg_end = float(seg.get("end"))
        except Exception:
            cleaned.append(seg)
            continue

        overlap_start = max(prev_start, seg_start)
        overlap_end = min(prev_end, seg_end)
        overlap = max(0.0, overlap_end - overlap_start)

        # Consider segments that overlap *or* directly touch with near-identical
        # text as duplicates to avoid repeated rows in downstream reports.
        touching = 0 <= (seg_start - prev_end) <= 0.25
        text_similarity = SequenceMatcher(
            None, (prev.get("text") or "").lower(), (seg.get("text") or "").lower()
        ).ratio()

        if not touching and overlap < min_overlap and text_similarity < 0.9:
            cleaned.append(seg)
            continue

        prev_other = _is_other_speaker(prev.get("speaker") or prev.get("speaker_name"))
        seg_other = _is_other_speaker(seg.get("speaker") or seg.get("speaker_name"))

        if prev_other and not seg_other:
            cleaned[-1] = seg
            continue
        if seg_other and not prev_other:
            continue

        prev_dur = _duration(prev)
        seg_dur = _duration(seg)
        prev_text = (prev.get("text") or "").strip()
        seg_text = (seg.get("text") or "").strip()

        if seg_dur > prev_dur or len(seg_text) > len(prev_text):
            cleaned[-1] = seg
            continue

        if prev_text and (prev_text == seg_text or text_similarity >= 0.9):
            prev["end"] = max(prev.get("end", overlap_end), seg.get("end", overlap_end))
            try:
                prev["duration"] = round(float(prev["end"]) - float(prev.get("start", 0.0)), 2)
            except Exception:
                prev.pop("duration", None)

    return cleaned


def diarization_from_audio(audio_url, transcript_segments, transcript_words=None):
    import contextlib
    import requests
    import subprocess
    import wave

    threshold = getattr(settings, "SPEAKER_MATCH_THRESHOLD", 0.8)
    min_embed_duration = getattr(settings, "SPEAKER_EMBEDDING_MIN_DURATION", 0.8)

    # Download audio locally for processing
    local_audio_path = os.path.join(tempfile.gettempdir(), f"diar_{os.path.basename(audio_url)}")
    with requests.get(audio_url, stream=True) as r:
        r.raise_for_status()
        with open(local_audio_path, 'wb') as f:
            for chunk in r.iter_content(chunk_size=8192):
                f.write(chunk)

    # Convert to mono 16k wav for consistency
    wav_path = local_audio_path
    if not local_audio_path.lower().endswith('.wav'):
        wav_path = local_audio_path.rsplit('.', 1)[0] + '.wav'
        command = [
            'ffmpeg', '-y', '-i', local_audio_path,
            '-ar', '16000', '-ac', '1', wav_path
        ]
        subprocess.run(command, check=True)
        os.unlink(local_audio_path)

    diarization = _get_diarization_pipeline()(wav_path)
    embedding_inference = _get_embedding_inference()

    audio_duration = None
    try:
        with contextlib.closing(wave.open(wav_path, "rb")) as wav_file:
            frame_rate = wav_file.getframerate() or 1
            audio_duration = wav_file.getnframes() / float(frame_rate)
    except Exception:
        logging.exception("Unable to determine audio duration for %s", wav_path)
        audio_duration = None
    
    def get_segment_text_from_words(words, seg_start, seg_end, overlap_threshold=0.1):
        """
        Extract text from words that overlap with the diarization segment.
        Uses overlap threshold to handle timing misalignments.
        """
        segment_words = []
        
        for word in words:
            word_start = word.get('start')
            word_end = word.get('end')
            
            if word_start is None or word_end is None:
                continue
            
            # Calculate overlap between word and diarization segment
            overlap_start = max(word_start, seg_start)
            overlap_end = min(word_end, seg_end)
            
            if overlap_end > overlap_start:
                # Calculate overlap percentage relative to word duration
                word_duration = word_end - word_start
                overlap_duration = overlap_end - overlap_start
                
                if word_duration > 0:
                    overlap_percentage = overlap_duration / word_duration
                    
                    # Include word if it has significant overlap
                    if overlap_percentage >= overlap_threshold:
                        segment_words.append({
                            'word': word.get('word', '').strip(),
                            'start': word_start,
                            'confidence': word.get('confidence', 1.0)
                        })
        
        # Sort by start time and join
        segment_words.sort(key=lambda x: x['start'])
        return " ".join(w['word'] for w in segment_words if w['word'])
    
    def get_segment_text_from_segments(segments, seg_start, seg_end, overlap_threshold=0.3):
        """
        Extract text from transcript segments that overlap with the diarization segment.
        """
        segment_texts = []
        
        for segment in segments:
            s_start = segment.get('start')
            s_end = segment.get('end')
            
            if s_start is None or s_end is None:
                continue
            
            # Calculate overlap
            overlap_start = max(s_start, seg_start)
            overlap_end = min(s_end, seg_end)
            
            if overlap_end > overlap_start:
                segment_duration = s_end - s_start
                overlap_duration = overlap_end - overlap_start
                
                if segment_duration > 0:
                    overlap_percentage = overlap_duration / segment_duration
                    
                    # Include segment if it has significant overlap
                    if overlap_percentage >= overlap_threshold:
                        segment_texts.append({
                            'text': segment.get('text', '').strip(),
                            'start': s_start
                        })
        
        # Sort by start time and join
        segment_texts.sort(key=lambda x: x['start'])
        return " ".join(t['text'] for t in segment_texts if t['text'])
    
    def merge_consecutive_same_speaker_segments(segments, max_gap=1.0):
        """Merge consecutive segments from the same speaker if they're close together."""
        if not segments:
            return segments

        merged = []

        def _init_segment(segment):
            seg = segment.copy()
            vectors = []
            vec = seg.get("speaker_vector")
            if vec:
                vectors.append(vec)
            seg["_vectors"] = vectors
            return seg

        current_segment = _init_segment(segments[0])

        for i in range(1, len(segments)):
            next_segment = _init_segment(segments[i])

            # Check if same speaker and segments are close
            if (
                current_segment['speaker'] == next_segment['speaker'] and
                next_segment['start'] - current_segment['end'] <= max_gap
            ):
                # Merge segments
                current_segment['end'] = next_segment['end']
                if next_segment['text'].strip():
                    if current_segment['text'].strip():
                        current_segment['text'] += " " + next_segment['text']
                    else:
                        current_segment['text'] = next_segment['text']
                current_segment["_vectors"].extend(next_segment.get("_vectors", []))
                if next_segment.get("duration"):
                    current_segment['duration'] = round(
                        current_segment.get('end', next_segment['end']) - current_segment.get('start', next_segment['start']), 2
                    )
            else:
                merged.append(current_segment)
                current_segment = next_segment

        merged.append(current_segment)

        # Compute averaged vectors and clean helper keys
        for seg in merged:
            vectors = seg.pop("_vectors", [])
            if vectors:
                try:
                    seg['speaker_vector'] = (
                        np.mean(np.array(vectors, dtype=float), axis=0).tolist()
                    )
                except Exception:
                    seg['speaker_vector'] = vectors[0]

        return merged
    
    # Process diarization results
    diarization_segments = []
    label_map: Dict[str, str] = {}
    label_index = 0

    for turn, _, speaker in diarization.itertracks(yield_label=True):
        seg_start = float(turn.start)
        seg_end = float(turn.end)

        # Extract text using the appropriate method
        if transcript_words:
            segment_text = get_segment_text_from_words(transcript_words, seg_start, seg_end)
        else:
            segment_text = get_segment_text_from_segments(transcript_segments, seg_start, seg_end)

        # Only add segments with actual content or significant duration
        if segment_text.strip() or (seg_end - seg_start) > 0.5:
            if speaker not in label_map:
                label_map[speaker] = f"SPEAKER_{label_index}"
                label_index += 1

            requested_segment = Segment(seg_start, seg_end)
            vector_list = None

            try:
                duration = seg_end - seg_start
                target_segment = requested_segment

                if duration < min_embed_duration and audio_duration:
                    center = seg_start + (duration / 2.0)
                    padded_start = max(0.0, center - min_embed_duration / 2.0)
                    padded_end = min(audio_duration, padded_start + min_embed_duration)

                    if padded_end - padded_start >= min_embed_duration * 0.5:
                        target_segment = Segment(padded_start, padded_end)

                if target_segment.duration > 0:
                    vector = embedding_inference.crop(wav_path, target_segment)
                    if vector is not None:
                        vector_list = vector.tolist()
            except Exception as exc:
                logging.warning(
                    "Skipping embedding for segment %.3f-%.3f: %s",
                    seg_start,
                    seg_end,
                    exc,
                )

            diarization_segments.append({
                "speaker": label_map[speaker],
                "speaker_label": label_map[speaker],
                "start": seg_start,
                "end": seg_end,
                "text": segment_text.strip(),
                "duration": round(seg_end - seg_start, 2),
                "speaker_vector": vector_list,
                "speaker_profile_id": None,
            })

    # Merge consecutive segments from the same speaker
    diarization_segments = merge_consecutive_same_speaker_segments(diarization_segments)

    # Filter out very short segments with no text
    diarization_segments = [
        seg for seg in diarization_segments 
        if seg['text'].strip() or seg['duration'] > 1.0
    ]
    
    # Attempt to match diarized speakers with stored profiles
    label_vectors: Dict[str, List[List[float]]] = {}
    for segment in diarization_segments:
        label = segment.get("speaker_label") or segment.get("speaker")
        vec = segment.get("speaker_vector")
        if label and vec:
            label_vectors.setdefault(label, []).append(vec)

    matched_profiles: Dict[str, SpeakerProfile] = {}
    label_means: Dict[str, List[float]] = {}
    for label, vectors in label_vectors.items():
        try:
            arr = np.array(vectors, dtype=float)
        except Exception:
            continue
        if arr.size == 0:
            continue
        if np.isnan(arr).any() or np.isinf(arr).any():
            continue
        mean_vec = np.mean(arr, axis=0)
        if mean_vec is None or np.isnan(mean_vec).any() or np.isinf(mean_vec).any():
            continue
        mean_vec_list = mean_vec.tolist()
        label_means[label] = mean_vec_list

        profile = match_speaker_embedding(mean_vec_list, threshold=threshold)
        if profile:
            matched_profiles[label] = profile

    # Persist new speaker profiles for unmatched speakers so they can be named later.
    for label, mean_vec in label_means.items():
        if label in matched_profiles:
            continue
        try:
            profile = SpeakerProfile.objects.create(
                embedding=mean_vec,
                name=label,
            )
        except Exception:
            continue
        matched_profiles[label] = profile

    for segment in diarization_segments:
        label = segment.get("speaker_label") or segment.get("speaker")
        profile = matched_profiles.get(label)
        if profile:
            segment["speaker"] = profile.name or label
            segment["speaker_name"] = profile.name
            segment["speaker_profile_id"] = profile.id
        else:
            segment.setdefault("speaker_name", None)

    # Remove overlapping placeholder segments that duplicate confident speakers.
    diarization_segments = _drop_overlapping_duplicates(diarization_segments)

    # Clean up
    if os.path.exists(wav_path):
        os.unlink(wav_path)

    return diarization_segments


# Diarization helpers

def build_speaker_summary(segments: Optional[List[Dict]]) -> List[Dict]:
    summary: Dict[str, Dict] = {}
    for seg in segments or []:
        label = seg.get('speaker_label') or seg.get('speaker')
        if not label:
            continue
        entry = summary.setdefault(
            label,
            {
                'speaker_label': label,
                'speaker_name': None,
                'speaker_profile_id': None,
                'segment_count': 0,
                'total_duration': 0.0,
            },
        )
        entry['segment_count'] += 1
        entry['total_duration'] += float(seg.get('duration') or 0.0)
        if seg.get('speaker_profile_id'):
            entry['speaker_profile_id'] = seg['speaker_profile_id']
            entry['speaker_name'] = seg.get('speaker_name') or seg.get('speaker')
        elif entry['speaker_name'] is None:
            entry['speaker_name'] = seg.get('speaker_name') or seg.get('speaker')

    for entry in summary.values():
        entry['total_duration'] = round(entry['total_duration'], 2)

    return list(summary.values())


def _format_timestamp(seconds: Optional[float]) -> str:
    if seconds is None:
        return "-"
    try:
        total_seconds = max(0, int(round(float(seconds))))
    except (TypeError, ValueError):
        return "-"

    hours, remainder = divmod(total_seconds, 3600)
    minutes, secs = divmod(remainder, 60)
    if hours:
        return f"{hours:02d}:{minutes:02d}:{secs:02d}"
    return f"{minutes:02d}:{secs:02d}"


def _format_time_range(start: Optional[float], end: Optional[float]) -> str:
    """Combine start and end timestamps into a compact range."""

    start_text = _format_timestamp(start)
    end_text = _format_timestamp(end)

    if start_text == "-" and end_text == "-":
        return "-"

    if start_text == "-":
        return end_text

    if end_text == "-":
        return start_text

    return f"{start_text} - {end_text}"

# Verification/Acknowledgment phrases (case-insensitive)
VERIFICATION_PHRASES = {
    "that's correct", "that's right", "correct","it is correct", "yes that's right",
    "that is correct", "yes that's correct", "yes correct", "you are correct",
}

READBACK_PHRASES = {
    "and you're", "you are", "so you", "you're ready", "proceed to", "moving to",
    "next step", "understand you", "i understand", "copy that", "you're done",
    "and you", "so you're", "you have", "you've", "you completed", "you are ready",
    "understand you", "understand it", "understand it's", "understand remove","understand section", "understand tube", "understand going", "understand we're",
    "understand we are", "understand the", "understand you're", "understand you are",
    "understand you want", "understand you need", "understand you need to",
}


def _is_verification_phrase(text: str, threshold: float = 0.80) -> bool:
    """Detect if text is a verification/acknowledgment phrase."""
    if not text:
        return False
    
    text_lower = text.lower().strip()
    word_count = len(text_lower.split())
    
    # Be more lenient with word count (up to 10 words)
    if word_count > 10:
        return False
    
    # Exact substring match (most reliable)
    for phrase in VERIFICATION_PHRASES:
        if phrase in text_lower:
            logging.debug(f"✓ Verification phrase detected: '{text_lower}' contains '{phrase}'")
            return True
    
    # Fuzzy match for short texts
    if word_count <= 6:
        for phrase in VERIFICATION_PHRASES:
            if fuzz.ratio(text_lower, phrase) >= threshold * 100:
                logging.debug(f"✓ Verification phrase detected (fuzzy): '{text_lower}' ≈ '{phrase}'")
                return True
    
    return False


def _is_readback_phrase(text: str) -> bool:
    """Detect if text contains readback language."""
    if not text:
        return False
    
    text_lower = text.lower()
    return any(phrase in text_lower for phrase in READBACK_PHRASES)


def _merge_incomplete_segments(segments: List[Dict], max_gap: float = 1.5) -> List[Dict]:
    """
    CONSERVATIVE merging - only merge truly fragmented segments.
    NEVER merge verification phrases or complete statements.
    """
    if not segments:
        return segments

    merged = []
    i = 0
    
    while i < len(segments):
        current = segments[i].copy()
        
        # Look ahead for potential merge
        while i + 1 < len(segments):
            next_seg = segments[i + 1]
            
            same_speaker = current.get('speaker') == next_seg.get('speaker')
            time_gap = next_seg['start'] - current['end']
            
            current_text = (current.get('text') or '').strip()
            next_text = (next_seg.get('text') or '').strip()
            
            # CRITICAL: Never merge verification phrases
            if _is_verification_phrase(current_text) or _is_verification_phrase(next_text):
                break
            
            # Check if incomplete (no period, short, or trailing connector)
            is_incomplete = (
                not current_text.endswith(('.', '!', '?')) or
                len(current_text.split()) < 3 or
                current_text.endswith((',', 'and', 'or', 'so', 'the', 'a', 'an', 'I', 'my'))
            )
            
            # Merge only if: same speaker, short gap, incomplete
            if same_speaker and time_gap <= max_gap and is_incomplete:
                current['end'] = next_seg['end']
                if current_text and next_text:
                    current['text'] = f"{current_text} {next_text}"
                elif next_text:
                    current['text'] = next_text
                current['duration'] = round(current['end'] - current['start'], 2)
                i += 1
            else:
                break
        
        merged.append(current)
        i += 1
    
    logging.info(f"Segment merging: {len(segments)} → {len(merged)}")
    return merged


def _split_segments_on_verification_phrases(segments: List[Dict]) -> List[Dict]:
    """Split segments so that verification phrases are isolated."""

    if not segments:
        return []

    # Build a single regex so that longer phrases are matched first and we can
    # capture every occurrence, even if multiple appear within the same
    # diarization segment.
    phrases = sorted(VERIFICATION_PHRASES, key=len, reverse=True)
    pattern = re.compile(r"(?i)\b(" + "|".join(re.escape(p) for p in phrases) + r")\b")
    trailing_punctuation = ",.;:!?)}]\"'”’"

    split_segments: List[Dict] = []

    for segment in segments:
        original_text = (segment.get("text") or "").strip()
        if not original_text:
            split_segments.append(segment)
            continue

        matches = list(pattern.finditer(original_text))
        if not matches:
            split_segments.append(segment)
            continue

        parts: List[str] = []
        cursor = 0
        for match in matches:
            start_idx, end_idx = match.span()

            # Leading speech before the verification phrase.
            leading = original_text[cursor:start_idx].strip()
            if leading:
                parts.append(leading)

            trailing_end = end_idx
            # Include trailing punctuation directly attached to the phrase.
            while (
                trailing_end < len(original_text)
                and original_text[trailing_end] in trailing_punctuation
            ):
                trailing_end += 1

            verification_text = original_text[start_idx:trailing_end].strip()
            parts.append(verification_text)
            cursor = trailing_end

        # Remainder after the last match.
        trailing_text = original_text[cursor:].strip()
        if trailing_text:
            parts.append(trailing_text)

        # Allocate timestamps proportionally across the new segments.
        start = segment.get("start")
        end = segment.get("end")
        total_duration = None
        if start is not None and end is not None:
            try:
                total_duration = max(float(end) - float(start), 0.0)
            except (TypeError, ValueError):
                total_duration = None

        weights = [max(len(text), 1) for text in parts]
        weight_sum = sum(weights) or 1
        current_start = float(start) if total_duration is not None else start

        for idx, text in enumerate(parts):
            new_segment = segment.copy()
            new_segment["text"] = text

            if total_duration is not None and start is not None and end is not None:
                portion = total_duration * (weights[idx] / weight_sum)
                new_start = current_start if current_start is not None else float(start)
                new_end = new_start + portion if idx < len(parts) - 1 else float(end)
                new_segment["start"] = round(new_start, 2)
                new_segment["end"] = round(new_end, 2)
                new_segment["duration"] = round(new_segment["end"] - new_segment["start"], 2)
                current_start = new_end
            else:
                new_segment["start"] = segment.get("start")
                new_segment["end"] = segment.get("end")
                if "duration" in new_segment and isinstance(new_segment.get("duration"), (int, float)):
                    new_segment["duration"] = segment.get("duration")

            split_segments.append(new_segment)

    logging.info(
        "Segment verification split: %d → %d", len(segments), len(split_segments)
    )
    return split_segments


def _combine_non_verification_runs(segments: List[Dict]) -> List[Dict]:
    """Combine adjacent non-verification segments spoken by the same speaker."""

    if not segments:
        return []

    combined: List[Dict] = []

    def _speaker_identity(segment: Dict) -> Optional[str]:
        return (
            segment.get("speaker")
            or segment.get("speaker_name")
            or segment.get("speaker_label")
        )

    for segment in segments:
        text = (segment.get("text") or "").strip()
        if not text:
            combined.append(segment)
            continue

        speaker_id = _speaker_identity(segment)
        if not combined:
            combined.append(segment.copy())
            continue

        previous = combined[-1]
        previous_text = (previous.get("text") or "").strip()
        previous_speaker = _speaker_identity(previous)

        if (
            speaker_id
            and speaker_id == previous_speaker
            and not _is_verification_phrase(previous_text)
            and not _is_verification_phrase(text)
        ):
            merged_text = " ".join(filter(None, [previous_text, text]))
            previous["text"] = merged_text.strip()

            # Update timing metadata when available.
            if segment.get("end") is not None:
                previous["end"] = segment.get("end")
            if previous.get("start") is not None and previous.get("end") is not None:
                try:
                    previous["duration"] = round(
                        float(previous["end"]) - float(previous["start"]), 2
                    )
                except (TypeError, ValueError):
                    previous.pop("duration", None)

            continue

        combined.append(segment.copy())

    logging.info(
        "Combined non-verification runs: %d → %d", len(segments), len(combined)
    )
    return combined


def _detect_3pc_patterns(segments: List[Dict]) -> List[Dict]:
    """
    Detect 3PC patterns: Statement → Readback → Confirmation
    IMPROVED: More aggressive verification detection
    """
    if len(segments) < 2:
        return segments
    
    for i in range(len(segments)):
        current = segments[i]
        current_text = (current.get('text') or '').strip()
        current_speaker = current.get('speaker')
        
        # PRIORITY: Look for verification phrases
        if _is_verification_phrase(current_text):
            if i > 0:
                prev = segments[i - 1]
                prev_speaker = prev.get('speaker')
                prev_text = (prev.get('text') or '').strip()
                
                # 3PC Confirmation: Different speaker + verification phrase
                if prev_speaker and prev_speaker != current_speaker:
                    current['_3pc_role'] = 'confirmation'
                    current['_3pc_confirms'] = i - 1
                    
                    # Mark previous as readback
                    prev['_3pc_role'] = 'readback'
                    
                    logging.info(f"✓ 3PC Confirmation: {current_speaker} confirms {prev_speaker}: '{current_text}'")
                    
                    # Look for original statement
                    if i > 1:
                        prev_prev = segments[i - 2]
                        if prev_prev.get('speaker') == current_speaker:
                            prev_prev['_3pc_role'] = 'statement'
                            prev['_3pc_confirms'] = i - 2
        
        # Look for readback language
        elif _is_readback_phrase(current_text) and i > 0:
            prev = segments[i - 1]
            if prev.get('speaker') != current_speaker:
                current['_3pc_role'] = 'readback'
                prev['_3pc_role'] = 'statement'
    
    return segments


def _assign_communication_groups(segments: List[Dict]) -> None:
    """Label segments as part of two-part or three-part communications."""

    group_id = 0
    processed_indices: Set[int] = set()

    for idx, segment in enumerate(segments):
        if segment.get("_3pc_role") != "confirmation" or idx in processed_indices:
            continue

        group_members: Set[int] = {idx}
        roles_present: Set[str] = {"confirmation"}

        readback_idx = segment.get("_3pc_confirms")
        statement_idx = None

        if isinstance(readback_idx, int) and 0 <= readback_idx < len(segments):
            readback_segment = segments[readback_idx]
            if readback_segment.get("_3pc_role") == "readback":
                group_members.add(readback_idx)
                roles_present.add("readback")
                statement_idx = readback_segment.get("_3pc_confirms")

        if isinstance(statement_idx, int) and 0 <= statement_idx < len(segments):
            statement_segment = segments[statement_idx]
            if statement_segment.get("_3pc_role") == "statement":
                group_members.add(statement_idx)
                roles_present.add("statement")

        # Fallback: confirmations may directly point to statements without readbacks.
        if ("readback" not in roles_present) and isinstance(readback_idx, int):
            statement_candidate = segments[readback_idx]
            if statement_candidate.get("_3pc_role") == "statement":
                group_members.add(readback_idx)
                roles_present.add("statement")

        if len(group_members) <= 1:
            continue

        group_id += 1
        communication_type = "3pc" if {"statement", "readback", "confirmation"} <= roles_present else "2pc"

        for member_idx in group_members:
            segments[member_idx]["_communication_group"] = group_id
            segments[member_idx]["_communication_type"] = communication_type
            processed_indices.add(member_idx)


def _create_document_buckets(
    reference_lines: List[str],
    model: SentenceTransformer,
    bucket_size: int = 3
) -> Tuple[List[Dict], Optional[List]]:
    """Create semantic buckets from reference document."""
    if not reference_lines:
        return [], None
    
    buckets = []
    for i in range(0, len(reference_lines), bucket_size):
        bucket_lines = reference_lines[i:i + bucket_size]
        bucket_text = " ".join(bucket_lines)
        buckets.append({
            "text": bucket_text,
            "lines": bucket_lines,
            "start_idx": i,
            "end_idx": min(i + bucket_size, len(reference_lines))
        })
    
    try:
        bucket_texts = [b["text"] for b in buckets]
        embeddings = model.encode(bucket_texts)
        return buckets, embeddings
    except Exception as e:
        logging.warning(f"Failed to create bucket embeddings: {e}")
        return buckets, None


def normalize_line(s: str) -> str:
    """Normalize text for comparison."""
    s = s.lower()
    s = re.sub(r"[\[\]\(\)\{\}\<\>]", "", s)
    s = s.translate(str.maketrans('', '', string.punctuation))
    return ' '.join(s.split())


def _find_best_bucket_match(
    spoken_text: str,
    buckets: List[Dict],
    bucket_embeddings: Optional[List],
    model: SentenceTransformer
) -> Tuple[Optional[int], float]:
    """Find best matching document bucket."""
    if not spoken_text or not buckets:
        return None, 0.0
    
    best_bucket_idx = None
    best_score = 0.0
    
    # Semantic matching
    if bucket_embeddings is not None:
        try:
            spoken_embedding = model.encode([spoken_text])
            similarities = util.cos_sim(spoken_embedding, bucket_embeddings)[0]
            best_bucket_idx = int(similarities.argmax())
            best_score = float(similarities[best_bucket_idx])
        except Exception:
            pass
    
    # Fuzzy matching fallback
    if best_score < 0.5:
        normalized_spoken = normalize_line(spoken_text)
        for idx, bucket in enumerate(buckets):
            normalized_bucket = normalize_line(bucket["text"])
            fuzzy_score = fuzz.token_set_ratio(normalized_spoken, normalized_bucket) / 100.0
            if fuzzy_score > best_score:
                best_score = fuzzy_score
                best_bucket_idx = idx

    return best_bucket_idx, best_score


def _similarity_label(value: Optional[float]) -> Optional[str]:
    if value is None:
        return None
    if value >= 0.7:
        return "High"
    if value >= 0.45:
        return "Medium"
    return "Low"


def _similarity_legend() -> str:
    return "Similarity Legend: High (≥0.70) | Medium (0.45–0.69) | Low (<0.45)"

def build_three_part_communication_summary(
    reference_text: Optional[str],
    diarization_segments: Optional[List[Dict]],
    match_threshold: float = 0.53,
    partial_threshold: float = 0.3,
    max_entries: int = 1000,
) -> List[Dict[str, Any]]:
    """
    CORRECTED Three-Part Communication (3PC) summary.
    
    FIXES:
    - "That's correct" now properly detected as "acknowledged"
    - 3PC pattern detection improved
    - All segments included (no truncation)
    """
    segments = diarization_segments or []

    # Clean up overlapping placeholder segments before further processing to
    # keep the 3PC report aligned with the final diarization output.
    segments = _drop_overlapping_duplicates(segments)
    
    if not segments:
        return []
    
    logging.info(f"Starting 3PC with {len(segments)} segments")
    
    # STEP 1: Conservative merging
    segments = _merge_incomplete_segments(segments, max_gap=1.5)

    # STEP 2: Split verification phrases that are bundled with other speech
    segments = _split_segments_on_verification_phrases(segments)
    # STEP 2a: Recombine the remaining speech for the same speaker
    segments = _combine_non_verification_runs(segments)

    # STEP 3: Detect 3PC patterns
    segments = _detect_3pc_patterns(segments)
    _assign_communication_groups(segments)

    # Parse reference document
    reference_lines: List[str] = []
    if reference_text:
        reference_lines = [ln.strip() for ln in reference_text.splitlines() if ln.strip()]
    
    # Get semantic model
    try:
        model = _get_sentence_model()
    except:
        model = SentenceTransformer('all-MiniLM-L6-v2', device = DEVICE)
    
    # Create document buckets
    buckets, bucket_embeddings = _create_document_buckets(reference_lines, model)

    summary_entries: List[Dict[str, Any]] = []
    used_bucket_indices: Set[int] = set()

    # Cache embeddings for repeated similarity checks between segment texts
    _text_embedding_cache: Dict[str, Any] = {}

    def _embed_text(text: str):
        key = (text or "").strip()
        if not key:
            return None
        if key not in _text_embedding_cache:
            _text_embedding_cache[key] = model.encode(key, convert_to_tensor=True)
        return _text_embedding_cache[key]

    def _similarity_to_text(text_a: str, text_b: str) -> float:
        emb_a = _embed_text(text_a)
        emb_b = _embed_text(text_b)
        if emb_a is None or emb_b is None:
            return 0.0
        try:
            sim_tensor = util.cos_sim(emb_a, emb_b)
            return float(sim_tensor.item())
        except Exception:
            return 0.0

    def _build_reason(
        status: str,
        role: Optional[str],
        matched_reference: Optional[str],
        similarity_value: Optional[float],
        context: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Construct a human-readable reason for the assigned status."""

        context = context or {}
        anchor_label = context.get("anchor_label")
        anchor_text = matched_reference or context.get("anchor_text")
        sim_label = _similarity_label(similarity_value)
        sim_text = f" ({sim_label} similarity)" if sim_label else ""
        note = context.get("note")

        if status == "match":
            if role == "readback":
                return f"Readback aligns with the {anchor_label or 'statement'}{sim_text}."
            if role == "confirmation":
                return "Acknowledgement confirms the prior instruction/readback."
            if role == "statement" and anchor_text:
                return f"Statement matches procedure text: \"{anchor_text}\"{sim_text}."
            if anchor_text:
                return f"Content matches procedure context{sim_text}."
            return "Content closely matches the expected text."

        if status == "partial match":
            if role == "readback":
                return f"Readback is partially aligned with the {anchor_label or 'statement'}{sim_text}."
            if anchor_text:
                return f"Partially matches procedure text{sim_text}."
            return "Content partially aligns with the expected text."

        if status == "mismatch":
            if role == "readback":
                return f"Readback differs from the {anchor_label or 'statement'}{sim_text}."
            if anchor_text:
                return f"No close match to procedure text{sim_text}."
            return "Content does not match the expected text."

        if status == "general conversation":
            base = "Short/non-procedural remark treated as general conversation."
            return f"{base} {note}".strip() if note else base

        if status == "acknowledged":
            return "Verification or acknowledgement phrase confirming prior instruction."

        if anchor_text and similarity_value is not None:
            return f"Compared against reference text{sim_text}."

        return "Status assigned based on dialogue context."

    def _classify_against_document(spoken: str):
        best_bucket_idx, best_score = _find_best_bucket_match(
            spoken, buckets, bucket_embeddings, model
        )

        matched_reference = None
        if best_bucket_idx is not None and best_score >= partial_threshold:
            bucket = buckets[best_bucket_idx]
            matched_reference = " ".join(bucket["lines"][:2])

        if best_score >= match_threshold:
            status = "match"
        elif best_score >= partial_threshold:
            status = "partial match"
        else:
            status = "mismatch"

        return status, matched_reference, best_score, best_bucket_idx

    def _normalize_partial_status(
        status: str,
        spoken: str,
        matched_reference: Optional[str],
        similarity_value: Optional[float],
        role: Optional[str] = None,
    ) -> str:
        if status != "partial match":
            return status

        word_count = len(spoken.split())
        low_confidence = (
            similarity_value is not None and similarity_value < (partial_threshold + 0.1)
        )
        missing_anchor = not matched_reference
        conversational_length = word_count < 4

        if role == "statement":
            if missing_anchor and (conversational_length or low_confidence):
                return "general conversation"
            return status

        if conversational_length or missing_anchor or low_confidence:
            return "general conversation"

        return status

    statement_context: Dict[int, Dict[str, Any]] = {}
    last_statement_idx: Optional[int] = None
    
    # STEP 3: Classify each segment
    for seg_idx, segment in enumerate(segments):
        spoken_text = (segment.get("text") or "").strip()
        speaker = (
            segment.get("speaker_name")
            or segment.get("speaker_label")
            or segment.get("speaker")
            or "Unknown"
        )
        
        if not spoken_text:
            continue
        
        three_pc_role = segment.get('_3pc_role')
        confirms_idx = segment.get('_3pc_confirms')
        
        # PRIORITY 1: 3PC Confirmation ("That's correct")
        if three_pc_role == 'confirmation':
            reference_content = None
            if confirms_idx is not None and confirms_idx < len(segments):
                confirmed_seg = segments[confirms_idx]
                reference_content = confirmed_seg.get('text', '')

            status = "acknowledged"
            similarity_label = _similarity_label(1.0)

            reason = _build_reason(
                status,
                "confirmation",
                reference_content,
                None,
                {
                    "anchor_label": "readback" if reference_content else None,
                    "anchor_text": reference_content,
                },
            )

            summary_entries.append({
                "speaker": speaker,
                "start": segment.get("start"),
                "end": segment.get("end"),
                "content": spoken_text,
                "status": status,
                "reference": reference_content,
                "similarity": similarity_label,
                "similarity_label": similarity_label,
                "similarity_score": 1.0,
                "three_pc_role": "confirmation",
                "communication_type": segment.get("_communication_type"),
                "reason": reason,
            })
            logging.info(f"✓ Added 3PC Confirmation: {speaker} - '{spoken_text}'")
            continue

        # PRIORITY 2: 3PC Readback
        if three_pc_role == 'readback':
            linked_statement_idx = None
            if seg_idx > 0 and segments[seg_idx - 1].get('_3pc_role') == 'statement':
                linked_statement_idx = seg_idx - 1
            if linked_statement_idx is None:
                linked_statement_idx = last_statement_idx

            statement_info = statement_context.get(linked_statement_idx)
            statement_text = (statement_info or {}).get("text")
            similarity_to_statement = _similarity_to_text(spoken_text, statement_text) if statement_text else 0.0

            if statement_text:
                if similarity_to_statement >= match_threshold:
                    status = "match"
                elif similarity_to_statement >= partial_threshold:
                    status = "partial match"
                else:
                    status = "mismatch"
                matched_reference = statement_text
                similarity_value = similarity_to_statement
            else:
                status, matched_reference, similarity_value, best_bucket_idx = _classify_against_document(spoken_text)
                if best_bucket_idx is not None and status == "match":
                    used_bucket_indices.add(best_bucket_idx)

            status = _normalize_partial_status(
                status,
                spoken_text,
                matched_reference,
                similarity_value,
                role="readback",
            )

            reason = _build_reason(
                status,
                "readback",
                matched_reference,
                similarity_value,
                {
                    "anchor_label": "statement" if statement_text else ("procedure" if matched_reference else None),
                    "anchor_text": statement_text or matched_reference,
                },
            )

            similarity_label = _similarity_label(similarity_value)
            
            summary_entries.append({
                "speaker": speaker,
                "start": segment.get("start"),
                "end": segment.get("end"),
                "content": spoken_text,
                "status": status,
                "reference": matched_reference,
                "similarity": similarity_label,
                "similarity_label": similarity_label,
                "similarity_score": round(similarity_value, 2) if similarity_value is not None else None,
                "three_pc_role": "readback",
                "communication_type": segment.get("_communication_type"),
                "reason": reason,
            })
            continue
        
        # PRIORITY 3: Standalone verification (not in 3PC)
        if _is_verification_phrase(spoken_text):
            reason = _build_reason(
                "acknowledged",
                "confirmation",
                None,
                1.0,
                {
                    "anchor_label": "statement" if last_statement_idx is not None else None,
                },
            )
            summary_entries.append({
                "speaker": speaker,
                "start": segment.get("start"),
                "end": segment.get("end"),
                "content": spoken_text,
                # Use the same key your color maps know (blue)
                "status": "acknowledged",
                "reference": None,
                "similarity": _similarity_label(1.0),
                "similarity_label": _similarity_label(1.0),
                "similarity_score": 1.0,
                # Make sure it renders with “(3PC: confirmation)”
                "three_pc_role": "confirmation",
                "communication_type": segment.get("_communication_type"),
                "reason": reason,
            })
            continue
        
        # STEP 4: Standard classification
        status, matched_reference, similarity_value, best_bucket_idx = _classify_against_document(spoken_text)

        downgrade_note = None

        # Statements must align to the procedure text; downgrade short utterances to mismatch
        if three_pc_role == 'statement':
            entry_status = status
        else:
            word_count = len(spoken_text.split())
            is_short = word_count < 4
            if best_bucket_idx is not None and status == "match":
                used_bucket_indices.add(best_bucket_idx)
            if not buckets:
                entry_status = "general conversation" if is_short else status
            else:
                if status == "mismatch" and is_short:
                    entry_status = "general conversation"
                else:
                    entry_status = status

        # Outside explicit 3PC exchanges, avoid over-confident matches that look conversational
        if (
            entry_status == "match"
            and three_pc_role is None
            and segment.get("_communication_type") is None
        ):
            similarity_label = _similarity_label(similarity_value)
            if similarity_label in {"Medium", "Low"}:
                if word_count <= 12 or not matched_reference:
                    entry_status = "general conversation"
                    downgrade_note = "Downgraded from match due to conversational context (medium/low similarity)."

        similarity_label = _similarity_label(similarity_value)
        entry_status = _normalize_partial_status(
            entry_status,
            spoken_text,
            matched_reference,
            similarity_value,
            role=three_pc_role,
        )

        reason = _build_reason(
            entry_status,
            three_pc_role,
            matched_reference,
            similarity_value,
            {
                "anchor_label": "procedure" if matched_reference else None,
                "anchor_text": matched_reference,
                "note": downgrade_note,
            },
        )

        entry = {
            "speaker": speaker,
            "start": segment.get("start"),
            "end": segment.get("end"),
            "content": spoken_text,
            "status": entry_status,
            "reference": matched_reference,
            "similarity": similarity_label,
            "similarity_label": similarity_label,
            "similarity_score": round(similarity_value, 2) if similarity_value is not None else None,
            "reason": reason,
        }

        if three_pc_role == 'statement':
            entry["three_pc_role"] = "statement"
            entry["communication_type"] = segment.get("_communication_type")
            statement_context[seg_idx] = {
                "text": spoken_text,
                "reference": matched_reference,
                "similarity": similarity_value,
            }
            last_statement_idx = seg_idx

        summary_entries.append(entry)
    
    # Sort by timestamp
    summary_entries.sort(key=lambda e: (
        0 if e.get("start") is not None else 1,
        float(e.get("start") or 0)
    ))
    
    # NO TRUNCATION - include all entries
    if len(summary_entries) > max_entries:
        logging.warning(f"3PC entries exceed {max_entries}, consider increasing limit")
    
    confirmations = sum(1 for e in summary_entries if e.get("three_pc_role") == "confirmation")
    logging.info(f"✓ 3PC Complete: {len(summary_entries)} entries, {confirmations} confirmations")
    
    return summary_entries
def append_three_pc_summary_to_pdf(pdf_path: str, entries: Optional[List[Dict]]) -> None:
    """
    Append 3PC summary to PDF with FIXED PAGINATION.
    Shows ALL entries across multiple pages.
    """
    import fitz
    import tempfile
    import textwrap
    import os
    
    doc = fitz.open(pdf_path)
    fd, temp_pdf_path = tempfile.mkstemp(suffix=".pdf")
    os.close(fd)
    saved = False

    try:
        margin = 36
        title_font_size = 16
        body_font_size = 10
        compact_font_size = 9
        row_padding = 4
        content_width_ratio = [0.16, 0.18, 0.32, 0.14, 0.2]
        headers = ["Speaker", "Time", "Content", "Status", "Reason"]
        column_font_sizes = [body_font_size, compact_font_size, body_font_size, body_font_size, compact_font_size]
        
        status_colors = {
            "match": (0, 0.5, 0),
            "acknowledged": (0, 0.4, 0.8),
            "partial match": (1, 0.65, 0),
            "general conversation": (0.5, 0.5, 0.5),
            "mismatch": (0.75, 0, 0),
        }

        def _new_page(include_header: bool = True):
            page = doc.new_page()
            width = page.rect.width
            height = page.rect.height
            y_pos = margin
            
            if include_header:
                page.insert_text(
                    (margin, y_pos),
                    "Three-Part Communication (3PC) Summary",
                    fontsize=title_font_size,
                    fontname="helv",
                )
                y_pos += title_font_size + 4

                legend_text = _similarity_legend()
                page.insert_text(
                    (margin, y_pos),
                    legend_text,
                    fontsize=compact_font_size,
                    fontname="helv",
                )
                y_pos += compact_font_size + 8

                content_width = width - 2 * margin
                x = margin
                for header, ratio in zip(headers, content_width_ratio):
                    col_width = content_width * ratio
                    rect = fitz.Rect(x, y_pos, x + col_width, y_pos + body_font_size + 8)
                    page.draw_rect(rect, color=(0, 0, 0))
                    page.insert_textbox(
                        rect, header,
                        fontsize=body_font_size,
                        fontname="helv",
                        align=fitz.TEXT_ALIGN_CENTER,
                    )
                    x += col_width
                y_pos += body_font_size + 8
            
            return page, y_pos, height

        def _wrap_cell(text: str, max_width: float, font_size: int) -> List[str]:
            if not text:
                return [""]
            approx_char_width = max(int(max_width / (font_size * 0.6)), 1)
            return textwrap.wrap(text, width=approx_char_width) or [""]

        def _draw_row(page, y_pos, page_height, values: List[str]):
            """Draw a row, creating new page if needed."""
            width = page.rect.width
            available_height = page_height - margin
            content_width = width - 2 * margin

            # Calculate row height
            columns = []
            max_lines = 1
            for value, ratio, font_size in zip(values, content_width_ratio, column_font_sizes):
                col_width = content_width * ratio
                lines = _wrap_cell(value, col_width - (2 * row_padding), font_size)
                columns.append((col_width, lines, font_size))
                max_lines = max(max_lines, len(lines))

            row_height = max_lines * (max(column_font_sizes) + 2) + (2 * row_padding)

            # CHECK: Do we need a new page?
            if y_pos + row_height > available_height:
                page, y_pos, page_height = _new_page(include_header=True)
                width = page.rect.width
                content_width = width - 2 * margin
                
                # Recalculate columns for new page
                columns = []
                max_lines = 1
                for value, ratio, font_size in zip(values, content_width_ratio, column_font_sizes):
                    col_width = content_width * ratio
                    lines = _wrap_cell(value, col_width - (2 * row_padding), font_size)
                    columns.append((col_width, lines, font_size))
                    max_lines = max(max_lines, len(lines))
                row_height = max_lines * (max(column_font_sizes) + 2) + (2 * row_padding)

            # Draw the row
            x = margin
            for idx, (col_width, lines, font_size) in enumerate(columns):
                rect = fitz.Rect(x, y_pos, x + col_width, y_pos + row_height)
                page.draw_rect(rect, color=(0.7, 0.7, 0.7))

                text_y = y_pos + row_padding + font_size
                color = (0, 0, 0)

                if headers[idx] == "Status":
                    status_key = values[idx].strip().lower()
                    base_status = status_key.split('(')[0].strip()
                    color = status_colors.get(base_status, color)

                for line in lines:
                    page.insert_text(
                        (x + row_padding, text_y),
                        line,
                        fontsize=font_size,
                        fontname="helv",
                        fill=color,
                    )
                    text_y += font_size + 2
                
                x += col_width

            return page, y_pos + row_height, page_height

        # Start first page
        page, y_cursor, page_height = _new_page()

        if not entries:
            page.insert_text(
                (margin, y_cursor),
                "No speaker diarization data available.",
                fontsize=body_font_size,
                fontname="helv",
            )
        else:
            # Draw ALL entries (fixed pagination)
            for idx, entry in enumerate(entries):
                status_text = (entry.get("status") or "").capitalize() or "-"
                if entry.get("three_pc_role"):
                    role = entry["three_pc_role"]
                    status_text += f" (3PC: {role})"

                row_values = [
                    entry.get("speaker") or "Unknown",
                    _format_time_range(entry.get("start"), entry.get("end")),
                    entry.get("content") or "",
                    status_text,
                    entry.get("reason") or "-",
                ]
                
                page, y_cursor, page_height = _draw_row(page, y_cursor, page_height, row_values)
                
                # Log progress every 50 entries
                if (idx + 1) % 50 == 0:
                    logging.info(f"PDF: Added {idx + 1}/{len(entries)} entries")

        doc.save(temp_pdf_path, deflate=True)
        saved = True
        logging.info(f"✓ 3PC PDF saved: {len(entries)} entries across {len(doc)} pages")
        
    finally:
        doc.close()
        if saved:
            os.replace(temp_pdf_path, pdf_path)
        if os.path.exists(temp_pdf_path):
            try:
                os.unlink(temp_pdf_path)
            except OSError:
                pass


def append_three_pc_summary_to_docx(docx_path: str, entries: Optional[List[Dict]]) -> None:
    """Append 3PC summary to DOCX with proper colors."""
    import docx
    from docx.shared import RGBColor, Pt, Inches
    
    document = docx.Document(docx_path)
    document.add_page_break()
    document.add_heading("Three-Part Communication (3PC) Summary", level=1)

    legend_para = document.add_paragraph(_similarity_legend())
    for run in legend_para.runs:
        run.font.size = Pt(9)
    legend_para.paragraph_format.space_after = Pt(6)

    if not entries:
        document.add_paragraph("No speaker diarization data available.")
        document.save(docx_path)
        return

    headers = ["Speaker", "Time", "Content", "Status", "Reason"]
    table = document.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        table.style = "Light Grid"

    table.autofit = False
    column_widths = [Inches(1.2), Inches(1.4), Inches(3.0), Inches(1.2), Inches(1.6)]

    for cell, title in zip(table.rows[0].cells, headers):
        cell.text = title

    status_colors = {
        "match": RGBColor(0, 128, 0),
        "acknowledged": RGBColor(0, 100, 200),
        "partial match": RGBColor(255, 165, 0),
        "general conversation": RGBColor(128, 128, 128),
        "mismatch": RGBColor(192, 0, 0),
    }

    for col, width in zip(table.columns, column_widths):
        for cell in col.cells:
            cell.width = width

    for entry in entries:
        row = table.add_row()
        cells = row.cells

        values = [
            entry.get("speaker") or "Unknown",
            _format_time_range(entry.get("start"), entry.get("end")),
            entry.get("content") or "",
        ]

        for idx, value in enumerate(values):
            cells[idx].text = value
            cells[idx].width = column_widths[idx]
            if headers[idx] in {"Time"}:
                for paragraph in cells[idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        status = (entry.get("status") or "").lower()
        status_cell = cells[len(headers) - 2]
        status_cell.width = column_widths[len(headers) - 2]
        
        status_text = status.capitalize() if status else "-"
        if entry.get("three_pc_role"):
            role = entry["three_pc_role"]
            status_text += f" (3PC: {role})"
        
        status_cell.text = status_text

        if status in status_colors:
            for paragraph in status_cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = status_colors[status]

        reason_cell = cells[len(headers) - 1]
        reason_cell.width = column_widths[len(headers) - 1]
        reason_cell.text = entry.get("reason") or "-"
        for paragraph in reason_cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

    document.save(docx_path)
    logging.info(f"✓ 3PC DOCX saved: {len(entries)} entries")

# --- CORE THREE-COLOR HIGHLIGHTING LOGIC ---

def _apply_color_to_paragraph_runs(p_element, color_hex_str):
    """Applies a color to all text runs within a paragraph's XML element."""
    for r_element in p_element.xpath('.//w:r'):
        rPr = r_element.find(qn('w:rPr'))
        if rPr is None:
            rPr = docx.oxml.OxmlElement('w:rPr')
            r_element.insert(0, rPr)
        
        color_element = docx.oxml.OxmlElement('w:color')
        color_element.set(qn('w:val'), color_hex_str)
        rPr.append(color_element)

def _collect_non_empty_paragraphs(element):
    if element is None:
        return []

    paragraphs = []
    for p_element in element.xpath('.//w:p'):
        full_text = "".join(p_element.xpath('.//w:t/text()')).strip()
        if not full_text:
            continue
        paragraphs.append((p_element, full_text, normalize_line(full_text)))
    return paragraphs


def _paragraph_index_for_offset(paragraphs, offset):
    """Map a character offset within the concatenated paragraph text to its index."""
    running = 0
    for idx, (_element, _text, normalized) in enumerate(paragraphs):
        end = running + len(normalized)
        if offset <= end:
            return idx
        running = end + 1  # account for the joining newline
    return max(len(paragraphs) - 1, 0)


def _detect_highlight_anchor_sliding(paragraphs, norm_trans, threshold):
    if not paragraphs or not norm_trans:
        return 0

    K = min(len(norm_trans), len(paragraphs), 3)
    if K == 0:
        return 0

    best_avg, best_idx = 0.0, 0
    for i in range(len(paragraphs) - K + 1):
        avg = sum(
            SequenceMatcher(None, paragraphs[i + j][2], norm_trans[j]).ratio()
            for j in range(K)
        ) / K
        if avg > best_avg:
            best_avg = avg
            best_idx = i

    return best_idx if best_avg >= threshold else 0


def _detect_highlight_anchor(paragraphs, norm_trans, threshold):
    if not paragraphs or not norm_trans:
        return 0

    doc_blob = "\n".join(p[2] for p in paragraphs if p[2])
    transcript_blob = "\n".join(t for t in norm_trans if t)

    if doc_blob and transcript_blob:
        matcher = SequenceMatcher(None, doc_blob, transcript_blob, autojunk=False)
        best_block = None
        best_quality = 0.0
        for block in matcher.get_matching_blocks():
            if not block.size:
                continue

            transcript_ratio = block.size / max(len(transcript_blob), 1)
            doc_ratio = block.size / max(len(doc_blob), 1)
            quality = max(transcript_ratio, doc_ratio)

            if block.size < 30 and transcript_ratio < threshold and doc_ratio < threshold:
                # Skip tiny overlaps unless they represent a sufficiently strong match.
                continue

            # Prefer longer, higher-quality matches to anchor the highlights.
            if quality > best_quality or (
                quality == best_quality and best_block and block.size > best_block.size
            ):
                best_quality = quality
                best_block = block

        if best_block and (best_quality >= threshold or best_block.size >= 80):
            return _paragraph_index_for_offset(paragraphs, best_block.a)

    # Fall back to the sliding window heuristic when we cannot locate a strong block
    return _detect_highlight_anchor_sliding(paragraphs, norm_trans, threshold)


def _apply_highlighting_to_paragraphs(paragraphs, norm_trans, thresholds, colors, start_index=0):
    if not paragraphs:
        return

    for idx, (p_element, _full_text, normalized_text) in enumerate(paragraphs):
        if idx < start_index:
            continue

        best_score = max((fuzz.token_set_ratio(normalized_text, t) for t in norm_trans), default=0) / 100.0

        color_to_apply = None
        if best_score >= thresholds['high']:
            color_to_apply = colors['GREEN']
        elif best_score >= thresholds['low']:
            color_to_apply = colors['RED']

        if color_to_apply:
            _apply_color_to_paragraph_runs(p_element, str(color_to_apply))


def _process_element_three_color(element, norm_trans, thresholds, colors, start_index=0):
    paragraphs = _collect_non_empty_paragraphs(element)
    _apply_highlighting_to_paragraphs(paragraphs, norm_trans, thresholds, colors, start_index)


def highlight_docx_three_color(docx_path, norm_trans, output_path, high_threshold=0.6, low_threshold=0.3):
    """Highlights text in a DOCX using the Green/Red/Black system."""
    document = docx.Document(docx_path)
    colors = {"GREEN": RGBColor(0, 176, 80), "RED": RGBColor(255, 0, 0)}
    thresholds = {'high': high_threshold, 'low': low_threshold}

    # Process main body, headers, and footers for complete coverage
    body_paragraphs = _collect_non_empty_paragraphs(document.element.body)
    body_start_index = _detect_highlight_anchor(body_paragraphs, norm_trans, thresholds['low'])
    _apply_highlighting_to_paragraphs(body_paragraphs, norm_trans, thresholds, colors, body_start_index)
    for section in document.sections:
        for part in [section.header, section.footer, section.first_page_header,
                     section.first_page_footer, section.even_page_header, section.even_page_footer]:
            _process_element_three_color(part._element, norm_trans, thresholds, colors)

    document.save(output_path)

def get_media_duration(file_obj):
    """
    Returns duration in seconds for an audio or video file object.
    Works on Windows by closing temp file before ffmpeg/moviepy reads it.
    """
    from moviepy import VideoFileClip, AudioFileClip

    # pick a reasonable suffix (extension) so ffmpeg can infer format
    suffix = os.path.splitext(getattr(file_obj, "name", ""))[-1] or ".media"

    tmp_path = None
    try:
        # Write to a *closed* temp file path (delete=False) so ffmpeg can open it
        fd, tmp_path = tempfile.mkstemp(suffix=suffix)
        os.close(fd)  # VERY IMPORTANT on Windows
        file_obj.seek(0)
        with open(tmp_path, "wb") as out:
            # If file_obj is large, copy in chunks
            chunk = file_obj.read(1024 * 1024)
            while chunk:
                out.write(chunk)
                chunk = file_obj.read(1024 * 1024)

        # Try as video first
        with suppress(Exception):
            with VideoFileClip(tmp_path) as clip:
                if clip.duration:
                    return float(clip.duration)

        # Then as audio
        with suppress(Exception):
            with AudioFileClip(tmp_path) as clip:
                if clip.duration:
                    return float(clip.duration)

        return 0.0
    except Exception:
        return 0.0
    finally:
        if tmp_path and os.path.exists(tmp_path):
            with suppress(Exception):
                os.remove(tmp_path)
